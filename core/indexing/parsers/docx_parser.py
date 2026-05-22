"""docx_parser.py — parse .docx files via python-docx."""

import hashlib
from datetime import datetime, timezone
from pathlib import Path

from .base import BaseParser


def _md5(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


class DocxParser(BaseParser):
    """Parse DOCX documents extracting paragraphs and tables."""

    def parse(self, filepath: Path) -> dict:
        from docx import Document

        doc = Document(str(filepath))

        # ---- paragraphs ----
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # ---- tables ----
        tables: list[list[str]] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    tables.append(cells)

        content = "\n\n".join(paragraphs)

        return {
            "source_file": filepath.name,
            "source_path": str(filepath),
            "content": content,
            "source_type": "docx",
            "char_count": len(content),
            "line_count": content.count("\n") + 1,
            "md5": _md5(content),
            "parsed_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "tables": tables,
        }
